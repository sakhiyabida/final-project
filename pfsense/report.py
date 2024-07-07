from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
import io

def add_header(document):
    try:
        header_section = document.sections[0]
        header = header_section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        logo_path = "static/Logo Horizontal.png"
        logo_height = Cm(1.5)
        run = header_paragraph.add_run()
        run.add_picture(logo_path, height=logo_height)
        header.add_paragraph()
    except Exception as e:
        print(f"Error in add_header: {e}")

def create_cover_page(document):
    try:
        add_header(document)
        document.add_picture('static/tri.png', width=Inches(2))
        document.paragraphs[-1].alignment = 1
        document.add_paragraph()
        document.paragraphs[-1].add_run().add_break()
        while len(document.paragraphs) % 4 != 0:
            document.add_paragraph()

        title = """Sample Penetration Test Report  
        Example Company"""
        document.add_heading(title, level=0).alignment = 1
        document.add_paragraph()
        document.paragraphs[-1].add_run().add_break()
        while len(document.paragraphs) % 14 != 0:
            document.add_paragraph()

        text = f"""
        Company: Customer Name
        Date: {date.today().strftime('%d %B %Y')}
        Version 1.0"""
        document.add_paragraph(text)
    except Exception as e:
        print(f"Error in create_cover_page: {e}")

def generate_findings(document):
    try:
        document.add_heading("Pendahuluan", level=2)
        intro_paragraph = document.add_paragraph()
        intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        intro_text = "Laporan ini dibuat sebagai hasil dari pengujian penetrasi yang dilakukan untuk mengidentifikasi dan mengekspos kerentanan pada sistem PfSense. Dalam laporan ini, kami akan memberikan detail mengenai temuan kerentanan yang kami identifikasi selama penilaian. Kami akan menjelaskan secara rinci potensi dampak dan risiko yang terkait dengan kerentanan ini, serta memberikan rekomendasi tindakan yang dapat diambil untuk memperbaiki kerentanan tersebut dan meningkatkan keamanan keseluruhan sistem."
        intro_paragraph.add_run(intro_text)

        document.add_heading("Ringkasan Eksekutif", level=2)
        summary_paragraph = document.add_paragraph()
        summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        summary_text = "Kami melakukan penetration testing pada tanggal " + date.today().strftime('%d %B %Y') + " dengan menggunakan kredensial atau pengetahuan sebelumnya tentang lingkungan internal. Tujuan pengujian ini adalah untuk mengidentifikasi kelemahan dan mencoba untuk mengeksploitasi kelemahan tersebut."
        summary_paragraph.add_run(summary_text)

        document.add_heading("Metodologi", level=2)
        methodology_paragraph = document.add_paragraph()
        methodology_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        methodology_text = "Pengujian penetrasi mengikuti metodologi komprehensif yang meliputi information gathering, vulnerability scanning, exploitation, dan post-exploitation."
        methodology_paragraph.add_run(methodology_text)

        document.add_heading("Temuan", level=2)
        findings_paragraph = document.add_paragraph()
        findings_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        findings_text = "Ketika kami melakukan penetration testing, kami menemukan bahwa di jaringan Anda terdapat kerentanan pada sistem PfSense."
        findings_paragraph.add_run(findings_text)

        document.add_heading("Pemindaian", level=2)
        with open("scan.txt", "r") as scan_file:
            scan_content = scan_file.read()
        document.add_paragraph(scan_content)

        document.add_heading("Eksploitasi", level=2)
        with open("poc.txt", "r") as poc_file:
            poc_content = poc_file.read()
        document.add_paragraph(poc_content)

        document.add_heading("Rekomendasi", level=2)
        recommendations_paragraph = document.add_paragraph()
        recommendations_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        recommendations_text = "Untuk mengurangi kerentanan tersebut, pengguna sebaiknya melakukan upgrade ke versi terbaru PfSense yang mencakup perbaikan untuk kerentanan tersebut. Selain itu, pengguna juga sebaiknya menerapkan praktik keamanan seperti berikut:"
        recommendations_paragraph.add_run(recommendations_text)

        recommendations_list = [
            "Melakukan pembaruan perangkat lunak secara teratur dan memastikan menggunakan versi terbaru PfSense.",
            "Memperkuat kontrol akses dengan menerapkan kebijakan otentikasi yang kuat dan penggunaan kata sandi yang kompleks",
            "Menonaktifkan akses root pada webGUI",
            "Melakukan pemantauan dan pencatatan aktivitas sistem secara teratur.",
            "Melakukan pelatihan kesadaran keamanan bagi anggota staf."
        ]
        for recommendation in recommendations_list:
            document.add_paragraph(recommendation, style='List Bullet')
    except Exception as e:
        print(f"Error in generate_findings: {e}")

def generate_report():
    try:
        document = Document()
        create_cover_page(document)
        generate_findings(document)

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        print("Penetration test report generated successfully.")
        return file_stream
    except Exception as e:
        print(f"Error generating report: {e}")
        return None

if __name__ == "__main__":
    report_stream = generate_report()
    if report_stream:
        with open("Penetration_Test_Report_PfSense.docx", "wb") as f:
            f.write(report_stream.read())
