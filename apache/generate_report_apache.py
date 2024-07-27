import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
import logging

# Set up logging
logging.basicConfig(
    filename='generate_report.log',
    level=logging.DEBUG,
    format='%(asctime)s %(levelname)s %(message)s'
)

def sanitize_text(text):
    # Replace non-XML-compatible characters with a placeholder
    return re.sub(r'[^\x09\x0A\x0D\x20-\x7E]', '?', text)

def add_header(document):
    try:
        # Add Header
        header_section = document.sections[0]
        header = header_section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Add the logo
        logo_path = "static/Logo Horizontal.png"  # Ensure the path is correct
        logo_height = Cm(1.5)
        run = header_paragraph.add_run()
        run.add_picture(logo_path, height=logo_height)
        # Add space after header
        header.add_paragraph()
        logging.info("Header added successfully.")
    except Exception as e:
        logging.error(f"Error adding header: {e}")

def create_cover_page(document):
    try:
        add_header(document)
        # Add the logo image
        document.add_picture('static/tri.png', width=Inches(2))  # Ensure the path is correct
        document.paragraphs[-1].alignment = 1  # Center align the image
        # Add spacer
        document.add_paragraph()
        document.paragraphs[-1].add_run().add_break()
        while len(document.paragraphs) % 4 != 0:
            document.add_paragraph()
        title = "Laporan Pengujian Penetrasi Server Apache HTTP"
        document.add_heading(title, level=0).alignment = 1  # Center align the text
        # Add spacer
        document.add_paragraph()
        document.paragraphs[-1].add_run().add_break()
        while len(document.paragraphs) % 14 != 0:
            document.add_paragraph()
        # Add details
        text = f"""
        Tanggal: {date.today().strftime('%d %B %Y')}
        Versi 1.0"""
        document.add_paragraph(text)
        logging.info("Cover page created successfully.")
    except Exception as e:
        logging.error(f"Error creating cover page: {e}")

def generate_report(scan_result, findings, ugm_logo_path, tri_logo_path):
    try:
        output_dir = os.path.join("static", "output")
        os.makedirs(output_dir, exist_ok=True)
        document = Document()

        create_cover_page(document)

        # Pendahuluan
        document.add_heading('Pendahuluan', level=1)
        introduction = (
            "Penelitian ini bertujuan untuk mengembangkan aplikasi berbasis web yang mengotomatisasi pengujian penetration testing terhadap kerentanan CVE-2021-42013 pada server Apache HTTP. "
            "Aplikasi ini dirancang untuk mempermudah dan mempercepat proses pengujian kerentanan, mulai dari pemindaian hingga pembuatan laporan. "
            "Hasil pengujian menunjukkan bahwa metode otomatis mampu menyelesaikan uji penetrasi dengan lebih cepat dan konsisten dibandingkan metode manual. "
            "Kerentanan yang ditemukan pada server Apache versi 2.4.50, terutama pada konfigurasi tertentu, memungkinkan eksploitasi path traversal dan remote code execution (RCE). "
            "Konfigurasi yang rentan termasuk mengaktifkan modul CGI-BIN, menambahkan direktori 'icons' dalam bagian Alias, dan mengatur `Require all granted` dalam konteks `<Directory>`. "
            "Konfigurasi ini memungkinkan serangan yang dapat membahayakan keamanan server dan jaringan."
        )
        paragraph = document.add_paragraph(introduction)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Findings
        document.add_heading('Hasil Temuan', level=1)
        document.add_heading('Hasil Pemindaian Nmap', level=2)
        document.add_paragraph(sanitize_text(scan_result))

        document.add_heading('Hasil Eksploitasi Kerentanan', level=2)
        document.add_paragraph(sanitize_text(findings))

        # Rekomendasi dan Mitigasi
        document.add_heading('Rekomendasi dan Mitigasi', level=1)
        recommendation = (
            "Jika Anda harus tetap menggunakan Apache versi 2.4.50 dan harus mengaktifkan modul CGI-BIN, berikut adalah beberapa langkah mitigasi yang dapat Anda ambil untuk mengurangi risiko kerentanan:\n"
            "1. Batasi Akses CGI-BIN:\n"
            "   - Hanya izinkan akses ke direktori CGI-BIN dari lokasi IP tertentu yang terpercaya.\n"
            "   ```\n"
            "   <Directory \"/path/to/cgi-bin\">\n"
            "       Options +ExecCGI\n"
            "       AllowOverride None\n"
            "       Order deny,allow\n"
            "       Deny from all\n"
            "       Allow from 192.168.1.0/24\n"
            "   </Directory>\n"
            "   ```\n"
            "2. Gunakan `ScriptAlias`:\n"
            "   - Pastikan direktori CGI-BIN hanya dapat diakses melalui `ScriptAlias` dan bukan melalui URL standar.\n"
            "   ```\n"
            "   ScriptAlias /cgi-bin/ \"/path/to/cgi-bin/\"\n"
            "   <Directory \"/path/to/cgi-bin\">\n"
            "       AllowOverride None\n"
            "       Options +ExecCGI -Indexes\n"
            "       Require all granted\n"
            "   </Directory>\n"
            "   ```\n"
            "3. Pemisahan Hak Istimewa:\n"
            "   - Jalankan skrip CGI dengan pengguna dan grup yang tidak memiliki hak istimewa tinggi untuk membatasi dampak jika ada skrip yang berhasil dieksploitasi.\n"
            "4. Monitoring dan Logging:\n"
            "   - Aktifkan logging terperinci untuk semua akses CGI.\n"
            "   ```\n"
            "   CustomLog ${APACHE_LOG_DIR}/cgi_access.log combined\n"
            "   ```\n"
            "   - Gunakan alat monitoring real-time seperti `ModSecurity` untuk mendeteksi dan merespons aktivitas mencurigakan.\n"
            "5. Pembatasan Akses File:\n"
            "   - Batasi akses ke direktori sensitif melalui konfigurasi Apache.\n"
            "   ```\n"
            "   <Directory \"/etc\">\n"
            "       Order deny,allow\n"
            "       Deny from all\n"
            "   </Directory>\n"
            "   ```\n"
            "6. Penggunaan WAF (Web Application Firewall):\n"
            "   - Pasang dan konfigurasi ModSecurity untuk melindungi server Apache dari berbagai jenis serangan, termasuk path traversal dan RCE.\n"
            "7. Penggunaan Proksi Terbalik:\n"
            "   - Gunakan reverse proxy seperti `nginx` atau `HAProxy` untuk menyaring permintaan sebelum mencapai server Apache, memberikan kontrol tambahan atas permintaan yang masuk.\n"
            "8. Audit dan Pengetesan Keamanan Berkala:\n"
            "   - Lakukan pengujian penetrasi secara berkala untuk menemukan dan memperbaiki potensi kerentanan.\n"
            "   - Lakukan audit keamanan berkala pada konfigurasi server dan skrip CGI yang digunakan.\n"
            "Dengan mengikuti langkah-langkah ini, risiko yang terkait dengan CVE-2021-42013 dan CVE-2021-41773 dapat dikurangi secara signifikan, membantu memastikan keamanan dan integritas Apache HTTP Server."
        )
        paragraph = document.add_paragraph(recommendation)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Save the report as a DOCX
        docx_report_path = os.path.join(output_dir, "Penetration_Test_Report_Apache.docx")
        document.save(docx_report_path)
        logging.info(f"DOCX report saved at {docx_report_path}")

        return docx_report_path
    except Exception as e:
        logging.error(f"Error generating report: {e}")
        return None
